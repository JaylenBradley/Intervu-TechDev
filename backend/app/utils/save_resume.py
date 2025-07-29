import re
from fpdf import FPDF
from docx import Document
from docx.shared import Pt

def save_text_as_pdf(text, output_path):
    BOLD_RE = re.compile(r'(\*\*.+\*\*)')
    sanitize_map = {
        '–': '-',  '—': '-',
        '“': '"',  '”': '"',
        '‘': "'",  '’': "'",
        '•': '-',  '…': '...',
    }
    for uni, ascii_char in sanitize_map.items():
        text = text.replace(uni, ascii_char)
    
    # Remove markdown bold markers and ensure consistent bullet points
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove ** ** bold markers
    text = text.replace('* ', '• ')  # Convert asterisk bullets to proper bullet points

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Times", size=10)

    lines = text.split('\n')
    first_line_processed = False
    
    for i, raw_line in enumerate(lines):
        line = raw_line.strip()
        if not line:
            pdf.ln(4)
            continue

        # Check if this is the first non-empty line (likely the name)
        if not first_line_processed and line and not line.isupper() and not line.startswith('*'):
            # This is likely the name - make it bold and larger
            pdf.set_font("Times", style='B', size=16)
            pdf.cell(0, 12, line)
            pdf.ln(12)
            pdf.set_font("Times", size=10)
            first_line_processed = True
            continue

        # HEADINGS
        if line.isupper():
            pdf.set_font("Times", style='B', size=11)
            pdf.cell(0, 8, line)
            pdf.ln(10)
            pdf.set_font("Times", size=10)
            continue

        # BULLETS
        if line.startswith('* ') or line.startswith('• '):
            bullet_text = line[2:].strip()
            pdf.cell(6)
            pdf.cell(4, 8, '•')
            parts = BOLD_RE.split(bullet_text)
            if len(parts) > 1:
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        pdf.set_font("Times", style='B', size=10)
                        pdf.write(8, part[2:-2])
                    else:
                        pdf.set_font("Times", size=10)
                        pdf.write(8, part)
            else:
                pdf.set_font("Times", size=10)
                pdf.write(8, bullet_text)
            pdf.ln(8)
            continue

        # INLINE BOLD
        parts = BOLD_RE.split(line)
        if len(parts) > 1:
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    pdf.set_font("Times", style='B', size=10)
                    pdf.write(8, part[2:-2])
                else:
                    pdf.set_font("Times", size=10)
                    pdf.write(8, part)
            pdf.ln(8)
            pdf.set_font("Times", size=10)
            continue

        # Plain text
        pdf.multi_cell(0, 8, line)

    pdf.output(output_path)

def save_text_as_docx(text, output_path):
    BOLD_RE = re.compile(r'(\*\*.+?\*\*)')
    sanitize_map = {
        '–': '-',  '—': '-',
        '“': '"',  '”': '"',
        '‘': "'",  '’': "'",
        '•': '-',  '…': '...',
    }
    for uni, ascii_char in sanitize_map.items():
        text = text.replace(uni, ascii_char)
    
    # Remove markdown bold markers and ensure consistent bullet points
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove ** ** bold markers
    text = text.replace('* ', '• ')  # Convert asterisk bullets to proper bullet points

    doc = Document()
    lines = text.split('\n')
    first_line_processed = False
    
    for i, raw_line in enumerate(lines):
        line = raw_line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Check if this is the first non-empty line (likely the name)
        if not first_line_processed and line and not line.isupper() and not line.startswith('*'):
            # This is likely the name - make it bold and larger
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(18)
            first_line_processed = True
            continue

        # HEADINGS
        if line.isupper():
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
            continue

        # BULLETS
        if line.startswith('* ') or line.startswith('• '):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parts = BOLD_RE.split(bullet_text)
            for part in parts:
                run = p.add_run(part[2:-2] if part.startswith('**') and part.endswith('**') else part)
                run.bold = part.startswith('**') and part.endswith('**')
            continue

        # INLINE BOLD
        parts = BOLD_RE.split(line)
        if len(parts) > 1:
            p = doc.add_paragraph()
            for part in parts:
                run = p.add_run(part[2:-2] if part.startswith('**') and part.endswith('**') else part)
                run.bold = part.startswith('**') and part.endswith('**')
            continue

        # Plain text
        doc.add_paragraph(line)

    doc.save(output_path)